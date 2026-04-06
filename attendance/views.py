from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from .models import Student, Subject, AttendanceRecord, AttendanceStatus, User
from django.contrib.auth.forms import UserCreationForm
import io
from datetime import datetime
from django.http import FileResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .forms import TeacherRegistrationForm
from django.utils.timezone import localtime
from django.contrib.auth import logout as django_logout


# ──────────────────────────────────────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────────────────────────────────────

def is_superuser(user):
    return user.is_superuser


def _pending_count(request):
    """Return pending teacher count for nav badge (used in context)."""
    if request.user.is_authenticated and request.user.is_superuser:
        return User.objects.filter(is_approved=False, is_superuser=False, is_active=True).count()
    return 0


# ──────────────────────────────────────────────────────────────────────────────
# AUTH / LANDING
# ──────────────────────────────────────────────────────────────────────────────

def index(request):
    if request.user.is_authenticated:
        if request.user.is_superuser or request.user.is_approved:
            return redirect('dashboard')
        return redirect('waiting_approval')
    return render(request, 'attendance/index.html')


@login_required
def waiting_approval(request):
    if request.user.is_superuser or request.user.is_approved:
        return redirect('dashboard')
    return render(request, 'attendance/waiting_approval.html')


@login_required
def renotify_admin(request):
    """Teacher presses 'Re-notify Admin' from waiting_approval page."""
    if request.method == 'POST':
        messages.success(request, "Admin has been notified again. Please wait for approval.")
    return redirect('waiting_approval')


def logout_view(request):
    django_logout(request)
    storage = messages.get_messages(request)
    for _ in storage:
        pass
    return redirect('index')


def register(request):
    if request.method == 'POST':
        form = TeacherRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.is_active = True
            user.is_approved = False
            user.save()
            messages.success(request, "Registration successful! Please wait for Admin approval.")
            return redirect('login')
    else:
        form = TeacherRegistrationForm()
    return render(request, 'attendance/register.html', {'form': form})


# ──────────────────────────────────────────────────────────────────────────────
# DASHBOARD
# ──────────────────────────────────────────────────────────────────────────────

@login_required
def dashboard(request):
    if not request.user.is_superuser and not request.user.is_approved:
        return redirect('waiting_approval')

    classes = Student.objects.values_list('student_class', flat=True).distinct().order_by('student_class')
    subjects = Subject.objects.all().order_by('code')

    if request.method == "POST":
        selected_class = request.POST.get('student_class')
        subject_id = request.POST.get('subject')
        return redirect('mark_attendance', s_class=selected_class, sub_id=subject_id)

    # Stats
    total_records  = AttendanceRecord.objects.count()
    total_students = Student.objects.count()
    recent_records = AttendanceRecord.objects.order_by('-date_created')[:6]

    # Pending teachers (superuser only)
    pending_teachers = []
    if request.user.is_superuser:
        pending_teachers = User.objects.filter(
            is_approved=False,
            is_superuser=False,
            is_active=True
        ).order_by('date_joined')

    return render(request, 'attendance/dashboard.html', {
        'classes':          classes,
        'subjects':         subjects,
        'is_admin':         request.user.is_superuser,
        'total_records':    total_records,
        'total_students':   total_students,
        'recent_records':   recent_records,
        'pending_teachers': pending_teachers,
        'pending_count':    len(pending_teachers),
    })


# ──────────────────────────────────────────────────────────────────────────────
# ATTENDANCE
# ──────────────────────────────────────────────────────────────────────────────

@login_required
def mark_attendance(request, s_class, sub_id):
    if not request.user.is_superuser and not request.user.is_approved:
        return redirect('waiting_approval')

    subject = get_object_or_404(Subject, id=sub_id)
    students = Student.objects.filter(student_class=s_class).order_by('university_roll')

    if request.method == "POST":
        record = AttendanceRecord.objects.create(
            student_class=s_class,
            subject=subject,
            teacher=request.user
        )

        present_student_ids = request.POST.getlist('present_students')
        entries = []
        for student in students:
            entries.append(
                AttendanceStatus(
                    record=record,
                    student=student,
                    is_present=str(student.id) in present_student_ids
                )
            )
        AttendanceStatus.objects.bulk_create(entries)
        messages.success(request, f"Attendance for {s_class} – {subject.name} saved successfully!")
        return redirect('dashboard')

    return render(request, 'attendance/mark_attendance.html', {
        'students': students,
        's_class':  s_class,
        'subject':  subject,
    })


# ──────────────────────────────────────────────────────────────────────────────
# REPORTS
# ──────────────────────────────────────────────────────────────────────────────

@login_required
def view_reports(request):
    if not request.user.is_superuser and not request.user.is_approved:
        return redirect('waiting_approval')
    records = AttendanceRecord.objects.all().order_by('-date_created')
    return render(request, 'attendance/reports.html', {
        'records':        records,
        'pending_count':  _pending_count(request),
    })


@login_required
def export_pdf(request, record_id):
    record   = get_object_or_404(AttendanceRecord, id=record_id)
    statuses = record.statuses.all().order_by('student__university_roll')

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    x_roll, x_name, x_status = 60, 220, 450
    row_height = 20

    def draw_main_header():
        p.setFont("Helvetica-Bold", 18)
        p.setFillColorRGB(0.1, 0.2, 0.5)
        p.drawCentredString(width / 2, 770, "Shri Ramswaroop Memorial College")
        p.setFont("Helvetica", 12)
        p.setFillColorRGB(0, 0, 0)
        p.drawCentredString(width / 2, 750, "Attendance Report")
        p.setFont("Helvetica-Bold", 12)
        p.drawCentredString(width / 2, 735, f"{record.subject.name} ({record.subject.code})")
        p.line(40, 720, 570, 720)
        p.setFont("Helvetica", 10)
        p.drawString(50, 700, f"Class: {record.student_class}")
        p.drawString(50, 685, f"Marked By: {record.teacher.get_full_name() or record.teacher.username}")
        p.drawString(50, 670, f"Attendance Time: {localtime(record.date_created).strftime('%d-%m-%Y %H:%M')}")
        p.drawString(50, 655, f"Report Generated At: {datetime.now().strftime('%d-%m-%Y %H:%M')}")

        total   = statuses.count()
        present = statuses.filter(is_present=True).count()
        absent  = total - present

        p.setFillColorRGB(0.95, 0.95, 0.95)
        p.roundRect(360, 655, 200, 60, 10, fill=1)
        p.setFont("Helvetica-Bold", 10)
        p.setFillColorRGB(0, 0, 0)
        p.drawString(370, 700, f"Total: {total}")
        p.setFillColorRGB(0, 0.6, 0)
        p.drawString(370, 685, f"Present: {present}")
        p.setFillColorRGB(0.8, 0, 0)
        p.drawString(370, 670, f"Absent: {absent}")

    def draw_table_header(y):
        p.setFillColorRGB(0.2, 0.2, 0.2)
        p.rect(50, y, 500, 25, fill=1)
        p.setFillColorRGB(1, 1, 1)
        p.setFont("Helvetica-Bold", 11)
        p.drawString(x_roll,   y + 7, "Roll No")
        p.drawString(x_name,   y + 7, "Student Name")
        p.drawString(x_status, y + 7, "Status")

    draw_main_header()
    y = 620
    draw_table_header(y)
    y -= 30

    for i, s in enumerate(statuses):
        if y < 60:
            p.showPage()
            y = 750
            draw_table_header(y)
            y -= 30

        if i % 2 == 0:
            p.setFillColorRGB(0.97, 0.97, 0.97)
            p.rect(50, y - 5, 500, row_height, fill=1)

        roll  = str(s.student.university_roll)
        name  = f"{s.student.first_name} {s.student.last_name}"
        base  = roll[:-3]
        last3 = roll[-3:]

        p.setFont("Helvetica", 10)
        p.setFillColorRGB(0, 0, 0)
        p.drawString(x_roll, y, base)
        w_base = p.stringWidth(base, "Helvetica", 10)

        p.setFont("Helvetica-Bold", 10)
        if s.is_present:
            p.setFillColorRGB(0, 0.6, 0)
        else:
            p.setFillColorRGB(0.8, 0, 0)
        p.drawString(x_roll + w_base, y, last3)

        p.setFont("Helvetica", 10)
        p.setFillColorRGB(0, 0, 0)
        p.drawString(x_name, y, name)

        p.setFont("Helvetica-Bold", 10)
        if s.is_present:
            p.setFillColorRGB(0, 0.6, 0)
            p.drawString(x_status, y, "PRESENT")
        else:
            p.setFillColorRGB(0.8, 0, 0)
            p.drawString(x_status, y, "ABSENT")

        p.setStrokeColorRGB(0.8, 0.8, 0.8)
        p.rect(50, y - 5, 500, row_height, fill=0)
        y -= row_height

    p.save()
    buffer.seek(0)
    return FileResponse(
        buffer,
        as_attachment=True,
        filename=f"Report_{record.student_class}_{record.date_created.date()}.pdf"
    )


@login_required
def export_docx(request, record_id):
    record   = get_object_or_404(AttendanceRecord, id=record_id)
    statuses = record.statuses.all().order_by('student__university_roll')

    doc = Document()

    title = doc.add_heading('Shri Ramswaroop Memorial College', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Attendance Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subject_line = doc.add_paragraph(f"{record.subject.name} ({record.subject.code})")
    subject_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "Class"
    table.cell(0, 1).text = record.student_class
    table.cell(1, 0).text = "Marked By"
    table.cell(1, 1).text = record.teacher.get_full_name() or record.teacher.username
    table.cell(2, 0).text = "Attendance Time"
    table.cell(2, 1).text = localtime(record.date_created).strftime('%d-%m-%Y %H:%M')
    table.cell(3, 0).text = "Report Generated At"
    table.cell(3, 1).text = datetime.now().strftime('%d-%m-%Y %H:%M')

    doc.add_paragraph("")

    total   = statuses.count()
    present = statuses.filter(is_present=True).count()
    absent  = total - present

    stats_para = doc.add_paragraph()
    stats_para.add_run(f"Total: {total}   ").bold = True
    run_p = stats_para.add_run(f"Present: {present}   ")
    run_p.bold = True
    run_p.font.color.rgb = RGBColor(0, 150, 0)
    run_a = stats_para.add_run(f"Absent: {absent}")
    run_a.bold = True
    run_a.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph("")

    s_table = doc.add_table(rows=1, cols=3)
    s_table.style = 'Table Grid'
    hdr = s_table.rows[0].cells
    hdr[0].text = "Roll No"
    hdr[1].text = "Name"
    hdr[2].text = "Status"

    for s in statuses:
        row  = s_table.add_row().cells
        roll = str(s.student.university_roll)
        base  = roll[:-3]
        last3 = roll[-3:]

        para  = row[0].paragraphs[0]
        run1  = para.add_run(base)
        run2  = para.add_run(last3)
        run2.bold = True
        if s.is_present:
            run2.font.color.rgb = RGBColor(0, 150, 0)
        else:
            run2.font.color.rgb = RGBColor(200, 0, 0)

        row[1].text = f"{s.student.first_name} {s.student.last_name}"

        para_s  = row[2].paragraphs[0]
        run_s   = para_s.add_run("PRESENT" if s.is_present else "ABSENT")
        run_s.bold = True
        if s.is_present:
            run_s.font.color.rgb = RGBColor(0, 150, 0)
        else:
            run_s.font.color.rgb = RGBColor(200, 0, 0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return FileResponse(
        buffer,
        as_attachment=True,
        filename=f"Report_{record.student_class}_{record.date_created.date()}.docx"
    )


# ──────────────────────────────────────────────────────────────────────────────
# STUDENT MANAGEMENT
# ──────────────────────────────────────────────────────────────────────────────

@login_required
@user_passes_test(is_superuser)
def manage_students(request):
    query        = request.GET.get('search', '')
    class_filter = request.GET.get('class_filter', '')
    all_classes  = Student.objects.values_list('student_class', flat=True).distinct().order_by('student_class')

    students = Student.objects.all().order_by('student_class', 'university_roll')

    if query:
        students = students.filter(university_roll__icontains=query) | \
                   Student.objects.filter(first_name__icontains=query) | \
                   Student.objects.filter(last_name__icontains=query)

    if class_filter:
        students = students.filter(student_class=class_filter)

    return render(request, 'attendance/manage_students.html', {
        'students':       students,
        'query':          query,
        'all_classes':    all_classes,
        'pending_count':  _pending_count(request),
    })


@login_required
@user_passes_test(is_superuser)
def delete_student(request, student_id):
    student = get_object_or_404(Student, id=student_id)
    if request.method == "POST":
        name = f"{student.first_name} {student.last_name}"
        student.delete()
        messages.success(request, f"Student '{name}' deleted successfully.")
    return redirect('manage_students')


# ──────────────────────────────────────────────────────────────────────────────
# TEACHER MANAGEMENT (SUPERUSER ONLY)
# ──────────────────────────────────────────────────────────────────────────────

@login_required
@user_passes_test(is_superuser)
def manage_teachers(request):
    tab              = request.GET.get('tab', 'pending')
    pending_teachers  = User.objects.filter(is_approved=False, is_superuser=False, is_active=True).order_by('date_joined')
    approved_teachers = User.objects.filter(is_approved=True,  is_superuser=False).order_by('date_joined')
    all_teachers      = User.objects.filter(is_superuser=False).order_by('-date_joined')

    return render(request, 'attendance/manage_teachers.html', {
        'tab':               tab,
        'pending_teachers':  pending_teachers,
        'approved_teachers': approved_teachers,
        'all_teachers':      all_teachers,
        'pending_count':     pending_teachers.count(),
    })


@login_required
@user_passes_test(is_superuser)
def approve_teacher(request, teacher_id):
    if request.method == 'POST':
        teacher = get_object_or_404(User, id=teacher_id)
        teacher.is_approved = True
        teacher.save()
        messages.success(request, f"'{teacher.get_full_name() or teacher.username}' has been approved.")
    return redirect(request.META.get('HTTP_REFERER', 'manage_teachers'))


@login_required
@user_passes_test(is_superuser)
def reject_teacher(request, teacher_id):
    """Reject (revoke) a teacher — sets is_approved=False. Superuser can also delete if needed."""
    if request.method == 'POST':
        teacher = get_object_or_404(User, id=teacher_id)
        teacher.is_approved = False
        teacher.save()
        messages.success(request, f"Access revoked for '{teacher.get_full_name() or teacher.username}'.")
    return redirect(request.META.get('HTTP_REFERER', 'manage_teachers'))